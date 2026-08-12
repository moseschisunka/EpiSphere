"""Report generation service"""

import csv
import re
import tempfile
from typing import Optional
from datetime import date, datetime
from sqlalchemy.orm import Session
from pathlib import Path

from app.db.models import Report, ReportType, Case, Country, Disease
from app.core.config import settings
from app.services.object_storage import PrivateObjectStorage


class ReportService:
    """Service for generating epidemiological report files."""

    def __init__(self, db: Session):
        self.db = db
        self.object_storage = PrivateObjectStorage()

    async def generate_report(
        self,
        report_type: ReportType,
        title: str,
        country_id: Optional[int] = None,
        disease_id: Optional[int] = None,
        start_date: Optional[date] = None,
        end_date: Optional[date] = None,
        file_format: str = "pdf",
        user_id: int = None
    ) -> Report:
        """Generate a report file and persist its metadata."""
        file_format = file_format.lower()
        if file_format not in {"pdf", "docx", "csv"}:
            raise ValueError("Unsupported report format")

        rows = self._get_case_rows(country_id, disease_id, start_date, end_date)
        summary = self._build_summary(title, report_type, rows, country_id, disease_id, start_date, end_date)
        object_key = self._build_object_key(title, file_format)
        with tempfile.TemporaryDirectory(prefix="episphere-report-") as temporary_dir:
            temporary_path = Path(temporary_dir) / f"report.{file_format}"
            if file_format == "pdf":
                self._write_pdf(temporary_path, summary, rows)
            elif file_format == "docx":
                self._write_docx(temporary_path, summary, rows)
            else:
                self._write_csv(temporary_path, summary, rows)
            stored_key = self.object_storage.store_file(
                temporary_path,
                object_key,
                self._content_type(file_format),
            )

        report = Report(
            title=title,
            report_type=report_type,
            country_id=country_id,
            disease_id=disease_id,
            start_date=start_date,
            end_date=end_date,
            file_path=stored_key,
            file_format=file_format,
            generated_by=user_id,
            report_metadata={
                "generated_at": datetime.utcnow().isoformat(),
                "report_type": report_type.value,
                "rows": len(rows),
                "total_cases": summary["total_cases"],
                "total_deaths": summary["total_deaths"],
                "object_storage_backend": self.object_storage.backend,
            },
        )

        self.db.add(report)
        self.db.commit()
        self.db.refresh(report)
        return report

    def _get_case_rows(self, country_id, disease_id, start_date, end_date):
        query = self.db.query(Case, Country, Disease).join(Country).join(Disease)
        if country_id:
            query = query.filter(Case.country_id == country_id)
        if disease_id:
            query = query.filter(Case.disease_id == disease_id)
        if start_date:
            query = query.filter(Case.date >= start_date)
        if end_date:
            query = query.filter(Case.date <= end_date)
        return query.order_by(Case.date.asc()).all()

    def _build_summary(self, title, report_type, rows, country_id, disease_id, start_date, end_date):
        total_cases = sum(row.Case.daily_cases for row in rows)
        total_deaths = sum(row.Case.daily_deaths for row in rows)
        return {
            "title": title,
            "report_type": report_type.value,
            "generated_at": datetime.utcnow().isoformat(),
            "country_id": country_id,
            "disease_id": disease_id,
            "start_date": start_date.isoformat() if start_date else None,
            "end_date": end_date.isoformat() if end_date else None,
            "row_count": len(rows),
            "total_cases": total_cases,
            "total_deaths": total_deaths,
            "cfr": round((total_deaths / total_cases) * 100, 2) if total_cases else None,
        }

    def _build_object_key(self, title: str, file_format: str) -> str:
        safe_title = re.sub(r"[^A-Za-z0-9_.-]+", "_", title).strip("_") or "report"
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        return f"{settings.REPORT_OBJECT_PREFIX.strip('/')}/{safe_title}_{timestamp}.{file_format}"

    @staticmethod
    def _content_type(file_format: str) -> str:
        return {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "csv": "text/csv",
        }[file_format]

    def _write_pdf(self, file_path: Path, summary: dict, rows):
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except ImportError as exc:
            raise RuntimeError("PDF generation requires reportlab") from exc

        pdf = canvas.Canvas(str(file_path), pagesize=letter)
        width, height = letter
        y = height - 48
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(48, y, summary["title"])
        y -= 28
        pdf.setFont("Helvetica", 10)
        for key in ["report_type", "generated_at", "row_count", "total_cases", "total_deaths", "cfr"]:
            pdf.drawString(48, y, f"{key.replace('_', ' ').title()}: {summary[key]}")
            y -= 16
        y -= 12
        pdf.setFont("Helvetica-Bold", 10)
        pdf.drawString(48, y, "Date")
        pdf.drawString(120, y, "Country")
        pdf.drawString(260, y, "Disease")
        pdf.drawString(390, y, "Cases")
        pdf.drawString(450, y, "Deaths")
        pdf.setFont("Helvetica", 9)
        y -= 14
        for row in rows[:35]:
            if y < 48:
                pdf.showPage()
                y = height - 48
                pdf.setFont("Helvetica", 9)
            pdf.drawString(48, y, row.Case.date.isoformat())
            pdf.drawString(120, y, row.Country.name[:24])
            pdf.drawString(260, y, row.Disease.name[:22])
            pdf.drawString(390, y, str(row.Case.daily_cases))
            pdf.drawString(450, y, str(row.Case.daily_deaths))
            y -= 14
        pdf.save()

    def _write_docx(self, file_path: Path, summary: dict, rows):
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX generation requires python-docx") from exc

        doc = Document()
        doc.add_heading(summary["title"], level=1)
        for key, value in summary.items():
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Date"
        hdr[1].text = "Country"
        hdr[2].text = "Disease"
        hdr[3].text = "Cases"
        hdr[4].text = "Deaths"
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = row.Case.date.isoformat()
            cells[1].text = row.Country.name
            cells[2].text = row.Disease.name
            cells[3].text = str(row.Case.daily_cases)
            cells[4].text = str(row.Case.daily_deaths)
        doc.save(str(file_path))

    def _write_csv(self, file_path: Path, summary: dict, rows):
        with file_path.open("w", newline="", encoding="utf-8") as fh:
            writer = csv.writer(fh)
            writer.writerow(["summary_key", "summary_value"])
            for key, value in summary.items():
                writer.writerow([key, value])
            writer.writerow([])
            writer.writerow(["date", "country", "disease", "daily_cases", "daily_deaths", "subnational_region", "source"])
            for row in rows:
                writer.writerow([
                    row.Case.date.isoformat(),
                    row.Country.name,
                    row.Disease.name,
                    row.Case.daily_cases,
                    row.Case.daily_deaths,
                    row.Case.subnational_region,
                    row.Case.source,
                ])
